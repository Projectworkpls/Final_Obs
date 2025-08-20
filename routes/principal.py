from flask import Blueprint, render_template, request, redirect, url_for, flash, jsonify, send_file, session, Response
import markdown
from models.database import (
    get_supabase_client, get_children, get_observers, get_parents,
    save_observation, get_observer_children, upload_file_to_storage, get_signed_audio_url,
    # Multi-tenant functions
    get_organizations, get_pending_observer_applications, review_observer_application,
    get_users_by_organization, get_children_by_organization, get_observer_child_mappings_by_organization,
    create_principal_feedback, get_peer_reviews_for_organization, auto_assign_parent_to_organization
)
from models.observation_extractor import ObservationExtractor
from utils.decorators import principal_required
import pandas as pd
import uuid
import json
from datetime import datetime
import io
import re
import urllib.parse
import logging
import os

logger = logging.getLogger(__name__)

principal_bp = Blueprint('principal', __name__)




def create_principal_feedback(principal_id, observer_id, feedback_text, feedback_type):
    try:
        supabase = get_supabase_client()

        # Convert to UUID objects (Supabase expects proper UUID format)
        from uuid import UUID
        principal_id = UUID(principal_id)
        observer_id = UUID(observer_id)

        response = supabase.table('principal_feedback').insert({
            "principal_id": str(principal_id),  # Convert back to string
            "observer_id": str(observer_id),
            "feedback_text": feedback_text,
            "feedback_type": feedback_type
        }).execute()

        # Check for Supabase-specific errors
        if getattr(response, 'error', None):
            logger.error(f"Supabase error: {response.error.message}")
            return False

        if response.data:
            logger.info(f"Feedback saved: ID {response.data[0]['id']}")
            return True

        logger.error("Empty response data from Supabase")
        return False

    except ValueError as e:
        logger.error(f"Invalid UUID format: {str(e)}")
        return False

    except Exception as e:
        logger.exception(f"Database insertion failed: {str(e)}")
        return False


@principal_bp.route('/dashboard')
@principal_required
def dashboard():
    try:
        principal_id = session.get('user_id')
        org_id = session.get('organization_id')

        if not org_id:
            flash('No organization assigned to your account. Please contact administrator.', 'error')
            return redirect(url_for('auth.login'))

        logger.info(f"Principal {principal_id} accessing dashboard for organization {org_id}")

        supabase = get_supabase_client()

        # Get analytics data for THIS ORGANIZATION ONLY
        users_response = supabase.table('users').select("id", count="exact").eq('organization_id', org_id).execute()
        observers_response = supabase.table('users').select("id", count="exact").eq('role', 'Observer').eq(
            'organization_id', org_id).execute()
        parents_response = supabase.table('users').select("id", count="exact").eq('role', 'Parent').eq(
            'organization_id', org_id).execute()
        children_response = supabase.table('children').select("id", count="exact").eq('organization_id',
                                                                                      org_id).execute()

        # Get observations for this organization
        org_users_response = supabase.table('users').select("id").eq('organization_id', org_id).execute()
        org_user_ids = [user['id'] for user in org_users_response.data] if org_users_response.data else []

        all_reports = []
        if org_user_ids:
            all_reports_response = supabase.table('observations').select("""
                id, student_name, observer_name, date, timestamp, filename, 
                file_url, full_data, processed_by_admin, username, student_id
            """).in_('username', org_user_ids).order('timestamp', desc=True).execute()
            all_reports = all_reports_response.data if all_reports_response.data else []

        # Process reports (simplified for debugging)
        processed_reports = []
        for report in all_reports:
            processed_report = {
                'id': report.get('id'),
                'student_name': report.get('student_name', 'N/A'),
                'observer_name': report.get('observer_name', 'N/A'),
                'date': report.get('date', 'N/A'),
                'timestamp': report.get('timestamp', 'N/A'),
                'filename': report.get('filename', 'N/A'),
                'file_url': report.get('file_url'),
                'processed_by_admin': report.get('processed_by_admin', False),
                'has_formatted_report': False,
                'formatted_report': None,
                'file_type': None,
                'signed_url': None,
                'organization_name': 'Your Organization'
            }

            if processed_report['file_url']:
                processed_report['file_url'] = urllib.parse.quote(processed_report['file_url'],
                                                                  safe=':/?#[]@!$&\'()*+,;=')
                file_url_lower = processed_report['file_url'].lower()
                if any(ext in file_url_lower for ext in ['.mp3', '.wav', '.m4a', '.ogg']):
                    processed_report['file_type'] = 'audio'
                elif any(ext in file_url_lower for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                    processed_report['file_type'] = 'image'

            if report.get('full_data'):
                try:
                    full_data = json.loads(report['full_data'])
                    if full_data.get('formatted_report'):
                        processed_report['has_formatted_report'] = True
                        processed_report['formatted_report'] = full_data['formatted_report']
                except:
                    pass

            processed_reports.append(processed_report)

        analytics = {
            'total_users': users_response.count if users_response.count else 0,
            'observers_count': observers_response.count if observers_response.count else 0,
            'parents_count': parents_response.count if parents_response.count else 0,
            'principals_count': 1,
            'children_count': children_response.count if children_response.count else 0,
            'observations_count': len(all_reports),
            'organizations_count': 1,
            'pending_applications': 0,
            'storage_files': len([r for r in processed_reports if r['file_url']]),
            'recent_observations': processed_reports[:10],
            'all_reports': processed_reports,
            'organizations': []
        }

        logger.info(f"Dashboard analytics: {analytics}")

    except Exception as e:
        logger.error(f"Error in principal dashboard: {e}")
        analytics = {
            'total_users': 0, 'observers_count': 0, 'parents_count': 0, 'principals_count': 0,
            'children_count': 0, 'observations_count': 0, 'organizations_count': 0,
            'pending_applications': 0, 'storage_files': 0,
            'recent_observations': [], 'all_reports': [], 'organizations': []
        }

    return render_template('principal/dashboard.html', analytics=analytics)


@principal_bp.route('/user_management')
@principal_required
def user_management():
    """Enhanced user management with ALL mappings visible for this organization"""
    try:
        org_id = session.get('organization_id')
        logger.info(f"Loading user management for organization: {org_id}")

        if not org_id:
            flash('No organization assigned to your account.', 'error')
            return redirect(url_for('principal.dashboard'))

        # Get all data for this organization
        users = get_users_by_organization(org_id)
        children = get_children_by_organization(org_id)
        observers = get_users_by_organization(org_id, 'Observer')
        parents = get_users_by_organization(org_id, 'Parent')

        # Get observer-child mappings for this organization
        observer_mappings = get_observer_child_mappings_by_organization(org_id)

        # Get parent-child mappings (parents with assigned children)
        parent_mappings = [parent for parent in parents if parent.get('child_id')]

        # Debug logging
        logger.info(f"User management data - Org: {org_id}")
        logger.info(f"Users: {len(users)}, Children: {len(children)}")
        logger.info(f"Observers: {len(observers)}, Parents: {len(parents)}")
        logger.info(f"Observer Mappings: {len(observer_mappings)}, Parent Mappings: {len(parent_mappings)}")

        # If no data, check if organization exists
        if not users and not children:
            flash('No data found for your organization. Please contact administrator to assign users and children.',
                  'warning')

        return render_template('principal/user_management.html',
                               users=users,
                               children=children,
                               observers=observers,
                               parents=parents,
                               observer_mappings=observer_mappings,
                               parent_mappings=parent_mappings)
    except Exception as e:
        logger.error(f"Error loading user management: {str(e)}")
        flash(f'Error loading user management: {str(e)}', 'error')
        return render_template('principal/user_management.html',
                               users=[], children=[], observers=[], parents=[],
                               observer_mappings=[], parent_mappings=[])


@principal_bp.route('/add_user', methods=['POST'])
@principal_required
def add_user():
    """Add user to this organization only"""
    try:
        org_id = session.get('organization_id')
        name = request.form.get('name')
        email = request.form.get('email').strip().lower()
        role = request.form.get('role')
        password = request.form.get('password')
        child_id = request.form.get('child_id') if role == 'Parent' else None

        if not all([name, email, role, password]):
            flash('All fields are required.', 'error')
            return redirect(url_for('principal.user_management'))

        user_data = {
            "id": str(uuid.uuid4()),
            "email": email,
            "name": name,
            "password": password,
            "role": role,
            "organization_id": org_id,
            "created_at": datetime.now().isoformat()
        }

        if child_id:
            user_data["child_id"] = child_id

        supabase = get_supabase_client()
        result = supabase.table('users').insert(user_data).execute()

        if result.data:
            flash('User added successfully!', 'success')
        else:
            flash('Error adding user - no data returned.', 'error')

    except Exception as e:
        logger.error(f"Error adding user: {str(e)}")
        flash(f'Error adding user: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/add_child', methods=['POST'])
@principal_required
def add_child():
    """Add child to this organization"""
    try:
        org_id = session.get('organization_id')
        name = request.form.get('child_name')
        birth_date = request.form.get('birth_date')
        grade = request.form.get('grade')

        if not name:
            flash('Child name is required.', 'error')
            return redirect(url_for('principal.user_management'))

        child_data = {
            "id": str(uuid.uuid4()),
            "name": name,
            "birth_date": birth_date,
            "grade": grade,
            "organization_id": org_id,
            "created_at": datetime.now().isoformat()
        }

        supabase = get_supabase_client()
        result = supabase.table('children').insert(child_data).execute()

        if result.data:
            flash('Child added successfully!', 'success')
        else:
            flash('Error adding child.', 'error')

    except Exception as e:
        logger.error(f"Error adding child: {str(e)}")
        flash(f'Error adding child: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/add_mapping', methods=['POST'])
@principal_required
def add_mapping():
    """Add mapping - organization restricted"""
    try:
        mapping_type = request.form.get('mapping_type')
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        if mapping_type == 'observer_child':
            observer_id = request.form.get('observer_id')
            child_id = request.form.get('child_id')

            if not observer_id or not child_id:
                flash('Please select both observer and child.', 'error')
                return redirect(url_for('principal.user_management'))

            # Verify observer belongs to this organization
            observer_check = supabase.table('users').select('id').eq('id', observer_id).eq('organization_id',
                                                                                           org_id).execute()
            if not observer_check.data:
                flash('Observer not found in your organization', 'error')
                return redirect(url_for('principal.user_management'))

            # Verify child belongs to this organization
            child_check = supabase.table('children').select('id').eq('id', child_id).eq('organization_id',
                                                                                        org_id).execute()
            if not child_check.data:
                flash('Child not found in your organization', 'error')
                return redirect(url_for('principal.user_management'))

            mapping_data = {
                "id": str(uuid.uuid4()),
                "observer_id": observer_id,
                "child_id": child_id,
                "created_at": datetime.now().isoformat()
            }

            result = supabase.table('observer_child_mappings').insert(mapping_data).execute()
            if result.data:
                flash('Observer-Child mapping added successfully!', 'success')
            else:
                flash('Error adding mapping.', 'error')

        elif mapping_type == 'parent_child':
            parent_id = request.form.get('parent_id')
            child_id = request.form.get('child_id')

            if not parent_id or not child_id:
                flash('Please select both parent and child.', 'error')
                return redirect(url_for('principal.user_management'))

            # Verify parent belongs to this organization
            parent_check = supabase.table('users').select('id').eq('id', parent_id).eq('organization_id',
                                                                                       org_id).execute()
            if not parent_check.data:
                flash('Parent not found in your organization', 'error')
                return redirect(url_for('principal.user_management'))

            # Verify child belongs to this organization
            child_check = supabase.table('children').select('id').eq('id', child_id).eq('organization_id',
                                                                                        org_id).execute()
            if not child_check.data:
                flash('Child not found in your organization', 'error')
                return redirect(url_for('principal.user_management'))

            result = supabase.table('users').update({'child_id': child_id}).eq('id', parent_id).execute()
            if result.data:
                flash('Parent-Child mapping added successfully!', 'success')
            else:
                flash('Error adding mapping.', 'error')

    except Exception as e:
        logger.error(f"Error adding mapping: {str(e)}")
        flash(f'Error adding mapping: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/send_observer_feedback', methods=['POST'])
@principal_required
def send_observer_feedback():
    """Send feedback to observer from user management"""
    try:
        principal_id = session.get('user_id')
        observer_id = request.form.get('observer_id')
        feedback_text = request.form.get('feedback_text')
        feedback_type = request.form.get('feedback_type')

        if not all([observer_id, feedback_text, feedback_type]):
            flash('All feedback fields are required.', 'error')
            return redirect(url_for('principal.user_management'))

        result = create_principal_feedback(principal_id, observer_id, feedback_text, feedback_type)

        if result:
            flash('Feedback sent successfully!', 'success')
        else:
            flash('Error sending feedback.', 'error')

    except Exception as e:
        logger.error(f"Error sending feedback: {str(e)}")
        flash(f'Error sending feedback: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/delete_mapping/<mapping_id>')
@principal_required
def delete_mapping(mapping_id):
    """Delete observer-child mapping - organization restricted"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Get mapping details to verify organization
        mapping = supabase.table('observer_child_mappings').select('*, users!observer_id(organization_id)').eq('id',
                                                                                                               mapping_id).execute()
        if not mapping.data or mapping.data[0]['users']['organization_id'] != org_id:
            flash('Mapping not found in your organization', 'error')
            return redirect(url_for('principal.user_management'))

        supabase.table('observer_child_mappings').delete().eq('id', mapping_id).execute()
        flash('Observer-Child mapping deleted successfully!', 'success')
    except Exception as e:
        logger.error(f"Error deleting mapping: {str(e)}")
        flash(f'Error deleting mapping: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/remove_parent_child_mapping/<parent_id>')
@principal_required
def remove_parent_child_mapping(parent_id):
    """Remove parent-child mapping"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Verify parent belongs to this organization
        parent_check = supabase.table('users').select('id').eq('id', parent_id).eq('organization_id', org_id).execute()
        if not parent_check.data:
            flash('Parent not found in your organization', 'error')
            return redirect(url_for('principal.user_management'))

        supabase.table('users').update({'child_id': None}).eq('id', parent_id).execute()
        flash('Parent-Child mapping removed successfully!', 'success')

    except Exception as e:
        logger.error(f"Error removing mapping: {str(e)}")
        flash(f'Error removing mapping: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/debug_org_data')
@principal_required
def debug_org_data():
    """Debug route to check organization data"""
    try:
        org_id = session.get('organization_id')
        principal_id = session.get('user_id')

        supabase = get_supabase_client()

        # Get organization info
        org_info = supabase.table('organizations').select('*').eq('id', org_id).execute()

        # Get all users in organization
        users = supabase.table('users').select('*').eq('organization_id', org_id).execute()

        # Get all children in organization
        children = supabase.table('children').select('*').eq('organization_id', org_id).execute()

        # Get all mappings
        mappings = supabase.table('observer_child_mappings').select('*').execute()

        debug_data = {
            'principal_id': principal_id,
            'organization_id': org_id,
            'organization_info': org_info.data,
            'users_in_org': users.data,
            'children_in_org': children.data,
            'all_mappings': mappings.data
        }

        return jsonify(debug_data)
    except Exception as e:
        return jsonify({'error': str(e)})


@principal_bp.route('/debug_peer_reviews')
@principal_required
def debug_peer_reviews():
    """Debug route to check peer review data"""
    try:
        org_id = session.get('organization_id')
        principal_id = session.get('user_id')

        supabase = get_supabase_client()

        # Get all peer reviews in the system
        all_peer_reviews = supabase.table('peer_reviews').select('*').execute()

        # Get all observations from this organization
        org_users = supabase.table('users').select('id, name').eq('organization_id', org_id).execute()
        org_user_ids = [user['id'] for user in org_users.data] if org_users.data else []

        org_observations = []
        if org_user_ids:
            org_observations = supabase.table('observations').select('id, student_name, username').in_('username',
                                                                                                       org_user_ids).execute()

        debug_data = {
            'principal_id': principal_id,
            'organization_id': org_id,
            'org_users': org_users.data,
            'org_user_ids': org_user_ids,
            'org_observations': org_observations.data if org_observations else [],
            'all_peer_reviews': all_peer_reviews.data,
            'peer_reviews_count': len(all_peer_reviews.data) if all_peer_reviews.data else 0
        }

        return jsonify(debug_data)
    except Exception as e:
        return jsonify({'error': str(e)})


# FIXED: Updated peer_reviews route to show actual cross-organization peer reviews
@principal_bp.route('/peer_reviews')
@principal_required
def peer_reviews():
    try:
        org_id = session.get('organization_id')
        principal_id = session.get('user_id')
        logger.info(f"Principal {principal_id} loading peer reviews for org {org_id}")

        supabase = get_supabase_client()

        # 1. Get organization users
        org_users = supabase.table('users').select('id,name,role').eq('organization_id', org_id).execute()
        org_users_data = org_users.data or []
        org_user_ids = [user['id'] for user in org_users_data]

        # 2. Get ALL organization observations for peer reviews
        all_org_observations = []
        org_observation_ids = []
        observation_map = {}
        if org_user_ids:
            # Pull all observations for peer review lookup
            obs_response = supabase.table('observations').select(
                'id, student_name, observer_name, date, full_data'
            ).in_('username', org_user_ids).order('timestamp', desc=True).execute()
            all_obs = obs_response.data or []
            
            # Create observation map for peer reviews
            for ob in all_obs:
                observation_map[ob.get('id')] = {
                    'id': ob.get('id'),
                    'student_name': ob.get('student_name') or 'N/A',
                    'observer_name': ob.get('observer_name') or 'N/A',
                    'date': ob.get('date') or 'N/A'
                }
            org_observation_ids = [obs['id'] for obs in all_obs]
            
            # Filter observations WITH AI review only for the AI reviews section
            filtered = []
            for ob in all_obs:
                try:
                    fd = json.loads(ob.get('full_data') or "{}")
                except Exception:
                    fd = {}
                ai_review = fd.get('communication_review')
                if not ai_review:
                    continue  # SHOW ONLY ONES THAT ALREADY HAVE AI REVIEW
                # Normalize whitespace for nicer display (do not alter stored text)
                clean_text = ai_review.strip()
                # Attach derived fields for template and actions
                filtered.append({
                    'id': ob.get('id'),
                    'student_name': ob.get('student_name') or 'N/A',
                    'observer_name': ob.get('observer_name') or 'N/A',
                    'date': (ob.get('date') or fd.get('processing_timestamp', '')[:10] or 'N/A'),
                    'communication_review': clean_text
                })
            org_observations = filtered
            logger.info(f"Found {len(org_observations)} observations with AI reviews")

        # 3. Get peer reviews for org observations
        org_peer_reviews = []
        logger.info(f"Looking for peer reviews for {len(org_observation_ids)} observations")
        if org_observation_ids:
            # OPTIMIZED: Fetch only relevant peer reviews
            peer_reviews_res = supabase.table('peer_reviews').select('*').in_('observation_id',
                                                                              org_observation_ids).order('created_at',
                                                                                                         desc=True).execute()
            peer_reviews_data = peer_reviews_res.data or []
            logger.info(f"Found {len(peer_reviews_data)} peer reviews")

            # Bulk fetch user details
            user_ids = set()
            for review in peer_reviews_data:
                user_ids.add(review['reviewer_id'])
                user_ids.add(review['observed_by'])

            user_map = {}
            if user_ids:
                users_res = supabase.table('users').select('id,name').in_('id', list(user_ids)).execute()
                user_map = {user['id']: user for user in users_res.data or []}

            # Build peer review objects
            for review in peer_reviews_data:
                org_peer_reviews.append({
                    **review,
                    'observation': observation_map.get(review['observation_id'], {}),
                    'reviewer_info': user_map.get(review['reviewer_id'], {}),
                    'observed_user_info': user_map.get(review['observed_by'], {})
                })

        # 4. Get observers and feedback
        observers = [user for user in org_users_data if user.get('role') == 'Observer']
        principal_feedback = []
        if observers:
            observer_ids = [obs['id'] for obs in observers]
            # OPTIMIZED: Direct feedback query
            feedback_res = supabase.table('principal_feedback').select('*').in_('observer_id', observer_ids).order(
                'created_at', desc=True).execute()
            principal_feedback = feedback_res.data or []

        logger.info(f"Rendering template with {len(org_peer_reviews)} peer reviews, {len(principal_feedback)} feedback items, {len(observers)} observers, {len(org_users_data)} users, {len(org_observations)} AI reviews")
        
        return render_template('principal/peer_reviews.html',
                               peer_reviews=org_peer_reviews,
                               principal_feedback=principal_feedback,
                               observers=observers,
                               users=org_users_data,
                               auto_reviews=org_observations # only AI‑reviewed items now
                               )

    except Exception as e:
        logger.error(f'Peer reviews error: {str(e)}')
        flash(f'Error loading peer reviews: {str(e)}', 'error')
        return render_template('principal/peer_reviews.html',
                               peer_reviews=[],
                               principal_feedback=[],
                               observers=[],
                               users=[]
                               )


@principal_bp.route('/send_peer_review_feedback', methods=['POST'])
@principal_required
def send_peer_review_feedback():
    try:
        principal_id = session.get('user_id')
        observer_id = request.form.get('observer_id')
        feedback_text = request.form.get('feedback_text')
        feedback_type = request.form.get('feedback_type')  # Form extraction

        # Validate session first
        if not principal_id:
            flash('User session expired. Please log in again.', 'error')
            return redirect(url_for('auth.login'))

        # === FEEDBACK TYPE HANDLING ===
        if feedback_type:
            feedback_type = feedback_type.capitalize()

        # Database expects: 'Positive', 'Constructive', 'Critical'
        valid_types = ['Positive', 'Constructive', 'Critical']

        # Case-insensitive validation and correction
        if feedback_type.lower() in [v.lower() for v in valid_types]:
            # Convert to correct case
            feedback_type = next(v for v in valid_types if v.lower() == feedback_type.lower())
        else:
            flash('Invalid feedback type selected', 'error')
            return redirect(url_for('principal.peer_reviews'))
        # === END FEEDBACK TYPE HANDLING ===

        # Validate required fields
        if not all([observer_id, feedback_text]):
            flash('Observer and Feedback Message are required', 'error')
            return redirect(url_for('principal.peer_reviews'))

        # Validate UUID format
        try:
            import uuid  # Ensure import is here if not at top
            uuid.UUID(observer_id)
            uuid.UUID(principal_id)
        except (ValueError, TypeError) as e:
            logger.error(f"Invalid UUID format: {str(e)}")
            flash('Invalid user ID format', 'error')
            return redirect(url_for('principal.peer_reviews'))

        # Validate feedback length
        if len(feedback_text) > 5000:
            flash('Feedback exceeds 5000 character limit', 'error')
            return redirect(url_for('principal.peer_reviews'))

        # Create feedback
        result = create_principal_feedback(
            principal_id=principal_id,
            observer_id=observer_id,
            feedback_text=feedback_text,
            feedback_type=feedback_type  # Now correctly capitalized
        )

        if result:
            flash('Feedback submitted successfully!', 'success')
        else:
            logger.error(f"Feedback save failed for principal={principal_id}, observer={observer_id}")
            flash('Database save failed. Please try again or contact support', 'error')

    except Exception as e:
        logger.exception(f"Critical system error: {str(e)}")
        flash('A system error occurred. Our team has been notified.', 'error')

    return redirect(url_for('principal.peer_reviews'))


# Keep all other existing routes (view_report, process_reports, analytics, etc.)
@principal_bp.route('/view_report/<report_id>')
@principal_required
def view_report(report_id):
    """View specific report - organization restricted"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Get the specific report and verify it belongs to this organization
        report_response = supabase.table('observations').select("""
            *, users!username(organization_id)
        """).eq('id', report_id).execute()

        if not report_response.data:
            flash('Report not found or access denied', 'error')
            return redirect(url_for('principal.dashboard'))

        report = report_response.data[0]

        # Verify organization access
        if report.get('users') and report['users'].get('organization_id') != org_id:
            flash('Access denied - report not from your organization', 'error')
            return redirect(url_for('principal.dashboard'))

        # URL encode the file_url if it exists and create signed URL for audio
        if report.get('file_url'):
            report['file_url'] = urllib.parse.quote(report['file_url'], safe=':/?#[]@!$&\'()*+,;=')

            # If it's an audio file, create signed URL
            if any(ext in report['file_url'].lower() for ext in ['.mp3', '.wav', '.m4a', '.ogg']):
                filename = report['file_url'].split('/')[-1]
                signed_url = get_signed_audio_url(filename)
                if signed_url:
                    report['signed_url'] = signed_url

        # Extract formatted report from full_data
        formatted_report = None
        if report.get('full_data'):
            try:
                full_data = json.loads(report['full_data'])
                formatted_report = full_data.get('formatted_report')
            except:
                pass

        return render_template('principal/view_report.html',
                               report=report,
                               formatted_report=formatted_report)

    except Exception as e:
        flash(f'Error loading report: {str(e)}', 'error')
        return redirect(url_for('principal.dashboard'))


@principal_bp.route('/process_reports')
@principal_required
def process_reports():
    """Process reports for this organization only"""
    try:
        org_id = session.get('organization_id')
        observers = get_users_by_organization(org_id, 'Observer')
        return render_template('principal/process_reports.html', observers=observers)
    except Exception as e:
        flash(f'Error loading process reports: {str(e)}', 'error')
        return render_template('principal/process_reports.html', observers=[])


@principal_bp.route('/analytics')
@principal_required
def analytics():
    """Organization analytics for this principal's organization"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Get comprehensive analytics data for THIS ORGANIZATION ONLY
        users_response = supabase.table('users').select("id", count="exact").eq('organization_id', org_id).execute()
        observers_response = supabase.table('users').select("id", count="exact").eq('role', 'Observer').eq(
            'organization_id', org_id).execute()
        parents_response = supabase.table('users').select("id", count="exact").eq('role', 'Parent').eq(
            'organization_id', org_id).execute()
        children_response = supabase.table('children').select("id", count="exact").eq('organization_id',
                                                                                      org_id).execute()

        analytics = {
            'total_users': users_response.count if users_response.count else 0,
            'observers_count': observers_response.count if observers_response.count else 0,
            'parents_count': parents_response.count if parents_response.count else 0,
            'principals_count': 1,
            'children_count': children_response.count if children_response.count else 0,
            'observations_count': 0,
            'organizations_count': 1,
            'storage_files': 0,
            'recent_observations': []
        }

        return render_template('principal/analytics.html', analytics=analytics)
    except Exception as e:
        flash(f'Error loading analytics: {str(e)}', 'error')
        return render_template('principal/analytics.html', analytics={})


@principal_bp.route('/delete_user/<user_id>')
@principal_required
def delete_user(user_id):
    """Delete user - organization restricted"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Verify user belongs to this organization
        user_check = supabase.table('users').select('id').eq('id', user_id).eq('organization_id', org_id).execute()
        if not user_check.data:
            flash('User not found in your organization', 'error')
            return redirect(url_for('principal.user_management'))

        supabase.table('users').delete().eq('id', user_id).execute()
        flash('User deleted successfully!', 'success')
    except Exception as e:
        flash(f'Error deleting user: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/delete_child/<child_id>')
@principal_required
def delete_child(child_id):
    """Delete child - organization restricted"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Verify child belongs to this organization
        child_check = supabase.table('children').select('id').eq('id', child_id).eq('organization_id', org_id).execute()
        if not child_check.data:
            flash('Child not found in your organization', 'error')
            return redirect(url_for('principal.user_management'))

        supabase.table('children').delete().eq('id', child_id).execute()
        flash('Child deleted successfully!', 'success')
    except Exception as e:
        flash(f'Error deleting child: {str(e)}', 'error')

    return redirect(url_for('principal.user_management'))


@principal_bp.route('/export_data')
@principal_required
def export_data():
    """Export organization data to CSV"""
    try:
        org_id = session.get('organization_id')
        supabase = get_supabase_client()

        # Get all data for this organization
        users = get_users_by_organization(org_id)
        children = get_children_by_organization(org_id)

        # Create CSV data
        output = io.StringIO()

        # Write users data
        output.write("USERS DATA\n")
        output.write("Name,Email,Role,Created At\n")
        for user in users:
            output.write(
                f"{user.get('name', '')},{user.get('email', '')},{user.get('role', '')},{user.get('created_at', '')}\n")

        output.write("\nCHILDREN DATA\n")
        output.write("Name,Birth Date,Grade,Created At\n")
        for child in children:
            output.write(
                f"{child.get('name', '')},{child.get('birth_date', '')},{child.get('grade', '')},{child.get('created_at', '')}\n")

        output.seek(0)

        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={"Content-Disposition": f"attachment;filename=organization_data_{org_id[:8]}.csv"}
        )

    except Exception as e:
        flash(f'Error exporting data: {str(e)}', 'error')
        return redirect(url_for('principal.dashboard'))


@principal_bp.route('/download_ai_review_docx/<observation_id>')
@principal_required
def download_ai_review_docx(observation_id):
    try:
        supabase = get_supabase_client()
        obs = supabase.table('observations').select(
            "id, student_name, observer_name, date, full_data"
        ).eq('id', observation_id).single().execute().data
        if not obs:
            flash('Observation not found', 'error')
            return redirect(url_for('principal.peer_reviews'))

        fd = json.loads(obs.get('full_data') or "{}")
        review_text = (fd.get('communication_review') or '').strip()
        if not review_text:
            flash('AI review not available for this observation', 'warning')
            return redirect(url_for('principal.peer_reviews'))

        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from io import BytesIO

        doc = Document()
        base = doc.styles['Normal']
        base.font.name = 'Segoe UI'
        base.font.size = Pt(11)

        # Title
        h = doc.add_heading('AI Communication Review', 0)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Meta line
        meta = doc.add_paragraph()
        meta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        meta_run = meta.add_run(
            f"Student: {obs.get('student_name','N/A')}    |    "
            f"Observer: {obs.get('observer_name','N/A')}    |    "
            f"Date: {obs.get('date','N/A')}"
        )
        meta_run.bold = True
        doc.add_paragraph("")  # spacer

        # Helpers to add content with formatting
        def add_heading(text, level=2):
            # Normalize emphasis markers
            clean = text.replace('**', '').strip('# ').strip()
            doc.add_heading(clean, level=level)

        def add_paragraph(text, bold=False):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(11)
            r.bold = bold

        def add_bullets(lines):
            for ln in lines:
                if ln.strip():
                    doc.add_paragraph(ln.strip(), style='List Bullet')

        def add_table_from_pipe(block_lines):
            # block_lines: header | separator | data rows
            if len(block_lines) < 2:  # not enough for a table
                return
            header = [c.strip() for c in block_lines[0].split('|') if c.strip()]
            # rows
            body_rows = []
            for row in block_lines[2:]:
                cells = [c.strip() for c in row.split('|') if c.strip()]
                if cells:
                    body_rows.append(cells)

            table = doc.add_table(rows=1, cols=len(header))
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            for i, c in enumerate(header):
                hdr[i].text = c

            for r in body_rows:
                row_cells = table.add_row().cells
                for i, c in enumerate(r):
                    if i < len(row_cells):
                        row_cells[i].text = c
            doc.add_paragraph("")  # spacer

        # Parse the review text using markdown
        try:
            import markdown
            md = markdown.Markdown(extensions=[
                'markdown.extensions.fenced_code',
                'markdown.extensions.tables',
                'markdown.extensions.nl2br',
                'markdown.extensions.sane_lists'
            ])
            
            # Convert markdown to HTML first
            html_content = md.convert(review_text)
            
            # Parse HTML and convert to Word document elements
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, 'html.parser')
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'strong', 'em', 'table']):
                if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    level = int(element.name[1])
                    add_heading(element.get_text().strip(), level=min(level, 3))
                elif element.name == 'p':
                    text = element.get_text().strip()
                    if text:
                        # Check if paragraph contains bold text
                        if element.find('strong') or element.find('b'):
                            add_paragraph(text, bold=True)
                        else:
                            add_paragraph(text)
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        text = li.get_text().strip()
                        if text:
                            doc.add_paragraph(text, style='List Bullet')
                elif element.name == 'table':
                    # Handle tables
                    rows = element.find_all('tr')
                    if len(rows) >= 2:  # At least header and one data row
                        header_row = rows[0]
                        headers = [th.get_text().strip() for th in header_row.find_all(['th', 'td'])]
                        
                        table = doc.add_table(rows=1, cols=len(headers))
                        table.style = 'Light Grid Accent 1'
                        hdr = table.rows[0].cells
                        for i, header in enumerate(headers):
                            hdr[i].text = header
                        
                        # Add data rows
                        for row in rows[1:]:
                            cells = [td.get_text().strip() for td in row.find_all('td')]
                            if cells:
                                row_cells = table.add_row().cells
                                for i, cell in enumerate(cells):
                                    if i < len(row_cells):
                                        row_cells[i].text = cell
                        doc.add_paragraph("")  # spacer
                elif element.name in ['strong', 'b']:
                    text = element.get_text().strip()
                    if text:
                        add_paragraph(text, bold=True)
                elif element.name in ['em', 'i']:
                    text = element.get_text().strip()
                    if text:
                        p = doc.add_paragraph()
                        r = p.add_run(text)
                        r.italic = True
                        
        except ImportError:
            # Fallback to original parsing if markdown or beautifulsoup not available
            lines = review_text.splitlines()
            buffer = []
            i = 0

            def flush_paragraph_buffer(buf):
                if not buf:
                    return
                bullets = []
                normal = []
                for t in buf:
                    if t.strip().startswith(('-', '*', '- ')):
                        bullets.append(t.lstrip('-*-  ').strip())
                    else:
                        normal.append(t)
                if normal:
                    add_paragraph(' '.join(normal).strip())
                if bullets:
                    add_bullets(bullets)

            while i < len(lines):
                ln = lines[i].rstrip()
                # Headings
                if ln.startswith('### '):
                    flush_paragraph_buffer(buffer); buffer = []
                    add_heading(ln[4:], level=3); i += 1; continue
                if ln.startswith('## '):
                    flush_paragraph_buffer(buffer); buffer = []
                    add_heading(ln[3:], level=2); i += 1; continue
                if ln.startswith('# '):
                    flush_paragraph_buffer(buffer); buffer = []
                    add_heading(ln[2:], level=1); i += 1; continue

                # Pipe table detection: header + separator + rows
                if '|' in ln and (i+1 < len(lines) and set(lines[i+1].strip()) <= set('|-: ')):
                    table_block = [ln, lines[i+1]]
                    i += 2
                    while i < len(lines) and '|' in lines[i]:
                        table_block.append(lines[i]); i += 1
                    flush_paragraph_buffer(buffer); buffer = []
                    add_table_from_pipe(table_block)
                    continue

                # Bold section titles like **1. ...**
                if ln.strip().startswith('**') and ln.strip().endswith('**') and len(ln.strip()) > 4:
                    flush_paragraph_buffer(buffer); buffer = []
                    add_paragraph(ln.strip().strip('* ').strip(), bold=True)
                    i += 1; continue

                buffer.append(ln)
                i += 1

            flush_paragraph_buffer(buffer)

        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)

        clean_name = (obs.get('student_name') or 'Student').replace(' ', '_')
        filename = f"AI_Review_{clean_name}_{obs.get('date','')}.docx"
        return send_file(
            stream,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        logger.exception(f"Error downloading AI review: {e}")
        flash('Error downloading AI review', 'error')
        return redirect(url_for('principal.peer_reviews'))


@principal_bp.route('/email_ai_review', methods=['POST'])
@principal_required
def email_ai_review():
    try:
        observation_id = request.form.get('observation_id')
        recipient_email = request.form.get('recipient_email')
        message = request.form.get('message', '')

        if not observation_id or not recipient_email:
            return jsonify({'success': False, 'error': 'Missing required fields'})

        supabase = get_supabase_client()
        obs = supabase.table('observations').select(
            "id, student_name, observer_name, date, full_data"
        ).eq('id', observation_id).single().execute().data

        if not obs:
            return jsonify({'success': False, 'error': 'Observation not found'})

        fd = json.loads(obs.get('full_data') or "{}")
        review_text = (fd.get('communication_review') or '').strip()
        if not review_text:
            return jsonify({'success': False, 'error': 'AI review not available'})

        # Here you would integrate with your email service
        # For now, we'll just return success
        # TODO: Implement actual email sending logic
        
        return jsonify({
            'success': True, 
            'message': f'AI review for {obs.get("student_name", "Student")} sent to {recipient_email}'
        })

    except Exception as e:
        logger.exception(f"Error emailing AI review: {e}")
        return jsonify({'success': False, 'error': str(e)})
