import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import calendar
from datetime import datetime
import io
import google.generativeai as genai
from config import Config
import re
import docx
from docx.shared import Inches, Pt
from io import BytesIO
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
from docx2pdf import convert

class MonthlyReportGenerator:
    def __init__(self, supabase_client):
        self.supabase = supabase_client
        # Initialize Gemini AI for report generation
        genai.configure(api_key=Config.GOOGLE_API_KEY)

    def get_month_data(self, child_id, year, month):
        """Fetch all observations for a specific child in a given month"""
        start_date = f"{year}-{month:02d}-01"
        if month == 12:
            end_date = f"{year + 1}-01-01"
        else:
            end_date = f"{year}-{month + 1:02d}-01"

        try:
            response = self.supabase.table('observations').select("*") \
                .eq("student_id", child_id) \
                .gte("date", start_date) \
                .lt("date", end_date) \
                .execute()
            return response.data
        except Exception as e:
            return []

    def get_goal_progress(self, child_id, year, month):
        """Get goal progress data for the specified month"""
        start_date = f"{year}-{month:02d}-01"
        if month == 12:
            end_date = f"{year + 1}-01-01"
        else:
            end_date = f"{year}-{month + 1:02d}-01"

        try:
            goals_response = self.supabase.table('goals').select("*") \
                .eq("child_id", child_id) \
                .execute()
            goals = goals_response.data

            goal_progress = []
            for goal in goals:
                alignments_response = self.supabase.table('goal_alignments').select("*") \
                    .eq("goal_id", goal['id']) \
                    .execute()
                alignments = alignments_response.data

                relevant_alignments = []
                for alignment in alignments:
                    report_response = self.supabase.table('observations').select("date") \
                        .eq("id", alignment['report_id']) \
                        .execute()

                    if report_response.data:
                        report_date = report_response.data[0]['date']
                        if start_date <= report_date < end_date:
                            relevant_alignments.append(alignment)

                if relevant_alignments:
                    avg_score = sum(a['alignment_score'] for a in relevant_alignments) / len(relevant_alignments)
                    progress_trend = [a['alignment_score'] for a in relevant_alignments]

                    goal_progress.append({
                        'goal_text': goal['goal_text'],
                        'avg_score': avg_score,
                        'progress_trend': progress_trend,
                        'num_observations': len(relevant_alignments),
                        'status': goal.get('status', 'active')
                    })

            return goal_progress
        except Exception as e:
            return []

    def get_strength_areas(self, observations):
        """Extract and count strength areas from observations"""
        strength_counts = {}

        for obs in observations:
            if obs.get('strengths'):
                try:
                    strengths = json.loads(obs['strengths']) if isinstance(obs['strengths'], str) else obs['strengths']
                    for strength in strengths:
                        strength_counts[strength] = strength_counts.get(strength, 0) + 1
                except:
                    pass

        return dict(sorted(strength_counts.items(), key=lambda x: x[1], reverse=True))

    def get_development_areas(self, observations):
        """Extract and count development areas from observations"""
        development_counts = {}

        for obs in observations:
            if obs.get('areas_of_development'):
                try:
                    areas = json.loads(obs['areas_of_development']) if isinstance(obs['areas_of_development'], str) else \
                        obs['areas_of_development']
                    for area in areas:
                        development_counts[area] = development_counts.get(area, 0) + 1
                except:
                    pass

        return dict(sorted(development_counts.items(), key=lambda x: x[1], reverse=True))

    def get_communication_skills(self, observations):
        """Extract communication skills metrics from observations"""
        comm_skills = {
            'confidence': [],
            'clarity': [],
            'participation': [],
            'sequencing': []
        }

        for obs in observations:
            try:
                full_data = json.loads(obs.get('full_data', '{}'))
                report = full_data.get('formatted_report', '')

                # Extract confidence level
                confidence_match = re.search(r'Confidence level:.*?(Strong|Moderate|Weak)', report, re.IGNORECASE)
                if confidence_match:
                    level = confidence_match.group(1).lower()
                    comm_skills['confidence'].append(level)

                # Extract clarity of thought
                clarity_match = re.search(r'Clarity of thought:.*?(Clear|Moderate|Unclear)', report, re.IGNORECASE)
                if clarity_match:
                    level = clarity_match.group(1).lower()
                    comm_skills['clarity'].append(level)

                # Extract participation
                participation_match = re.search(r'Participation.*?:.*?(Active|Moderate|Passive)', report, re.IGNORECASE)
                if participation_match:
                    level = participation_match.group(1).lower()
                    comm_skills['participation'].append(level)

                # Extract sequence of explanation
                sequence_match = re.search(r'Sequence of explanation:.*?(Logical|Moderate|Disorganized)', report,
                                           re.IGNORECASE)
                if sequence_match:
                    level = sequence_match.group(1).lower()
                    comm_skills['sequencing'].append(level)

            except Exception as e:
                continue

        # Calculate most common values
        comm_summary = {}
        for skill, values in comm_skills.items():
            if values:
                comm_summary[skill] = max(set(values), key=values.count)
            else:
                comm_summary[skill] = 'no data'

        return comm_summary

    def get_growth_metrics(self, observations):
        """Extract growth metrics from observations"""
        growth_areas = {
            'Intellectual': [],
            'Emotional': [],
            'Social': [],
            'Creativity': [],
            'Physical': [],
            'Character/Values': [],
            'Planning/Independence': []
        }

        for obs in observations:
            try:
                full_data = json.loads(obs.get('full_data', '{}'))
                report = full_data.get('formatted_report', '')

                # Extract growth area ratings
                for area in growth_areas.keys():
                    pattern = re.compile(rf'{area}.*?\|.*?(\w+)\s*\|', re.IGNORECASE)
                    match = pattern.search(report)
                    if match:
                        rating = match.group(1).lower()
                        growth_areas[area].append(rating)

            except Exception as e:
                continue

        # Calculate most common ratings
        growth_summary = {}
        for area, ratings in growth_areas.items():
            if ratings:
                growth_summary[area] = max(set(ratings), key=ratings.count)
            else:
                growth_summary[area] = 'no data'

        return growth_summary

    def generate_observation_frequency_chart(self, observations):
        """Generate a chart showing the frequency of observations by date"""
        date_counts = {}

        for obs in observations:
            date = obs.get('date', '')
            if date:
                date_counts[date] = date_counts.get(date, 0) + 1

        if not date_counts:
            return None

        df = pd.DataFrame([
            {"date": date, "count": count}
            for date, count in date_counts.items()
        ])

        df['date'] = pd.to_datetime(df['date'])
        df = df.sort_values('date')

        fig = px.bar(
            df,
            x='date',
            y='count',
            title='📅 Observation Frequency by Date',
            labels={'date': 'Date', 'count': 'Number of Observations'}
        )

        return fig

    def generate_strengths_chart(self, strength_counts):
        """Generate a chart showing the frequency of different strengths"""
        if not strength_counts:
            return None

        top_strengths = dict(list(strength_counts.items())[:10])

        df = pd.DataFrame([
            {"strength": strength, "count": count}
            for strength, count in top_strengths.items()
        ])

        fig = px.bar(
            df,
            x='count',
            y='strength',
            title='🌟 Top Strengths Observed',
            labels={'strength': 'Strength', 'count': 'Frequency'},
            orientation='h'
        )

        return fig

    def generate_development_areas_chart(self, development_counts):
        """Generate a chart showing the frequency of different development areas"""
        if not development_counts:
            return None

        top_areas = dict(list(development_counts.items())[:10])

        df = pd.DataFrame([
            {"area": area, "count": count}
            for area, count in top_areas.items()
        ])

        fig = px.bar(
            df,
            x='count',
            y='area',
            title='📈 Areas for Development',
            labels={'area': 'Development Area', 'count': 'Frequency'},
            orientation='h'
        )

        return fig

    def generate_goal_progress_chart(self, goal_progress):
        """Generate a chart showing progress on goals"""
        if not goal_progress:
            return None

        fig = make_subplots(rows=len(goal_progress), cols=1,
                            subplot_titles=[g['goal_text'][:50] + '...' for g in goal_progress],
                            vertical_spacing=0.1)

        for i, goal in enumerate(goal_progress):
            fig.add_trace(
                go.Bar(
                    x=[goal['avg_score']],
                    y=['Average Score'],
                    orientation='h',
                    name=f"Goal {i + 1}",
                    showlegend=False
                ),
                row=i + 1, col=1
            )

            fig.add_shape(
                type="line",
                x0=10, y0=-0.5,
                x1=10, y1=0.5,
                line=dict(color="green", width=2, dash="dash"),
                row=i + 1, col=1
            )

        fig.update_layout(
            title_text="🎯 Goal Progress",
            height=200 * len(goal_progress),
            margin=dict(l=0, r=0, t=50, b=0)
        )

        return fig

    def generate_monthly_summary(self, observations, goal_progress):
        """Generate a text summary of the monthly progress"""
        if not observations:
            return "No observations recorded this month."

        num_observations = len(observations)
        num_goals_with_progress = len(goal_progress)

        if goal_progress:
            avg_goal_score = sum(g['avg_score'] for g in goal_progress) / len(goal_progress)
            highest_goal = max(goal_progress, key=lambda x: x['avg_score'])
            lowest_goal = min(goal_progress, key=lambda x: x['avg_score'])
        else:
            avg_goal_score = 0
            highest_goal = None
            lowest_goal = None

        summary = f"""
        ### 📋 Monthly Progress Summary

        **📊 Total Observations:** {num_observations}
        **🎯 Goals Tracked:** {num_goals_with_progress}
        **📈 Average Goal Progress:** {avg_goal_score:.1f}/10
        """

        if highest_goal:
            summary += f"""
            **🌟 Strongest Goal Area:** {highest_goal['goal_text'][:50]}... (Score: {highest_goal['avg_score']:.1f}/10)
            **📉 Goal Needing Most Support:** {lowest_goal['goal_text'][:50]}... (Score: {lowest_goal['avg_score']:.1f}/10)
            """

        return summary

    def generate_monthly_summary_json_format(self, observations, goal_progress, child_name, year, month):
        """Generate monthly summary in the new JSON format with graph recommendations"""
        try:
            # Get communication skills and growth metrics
            comm_skills = self.get_communication_skills(observations)
            growth_metrics = self.get_growth_metrics(observations)

            # Prepare data for analysis
            observation_texts = []
            all_strengths = []
            all_developments = []
            all_recommendations = []

            for obs in observations:
                observation_texts.append(obs.get('observations', ''))

                # Parse strengths, developments, and recommendations
                if obs.get('strengths'):
                    try:
                        strengths = json.loads(obs['strengths']) if isinstance(obs['strengths'], str) else obs[
                            'strengths']
                        all_strengths.extend(strengths)
                    except:
                        pass

                if obs.get('areas_of_development'):
                    try:
                        developments = json.loads(obs['areas_of_development']) if isinstance(
                            obs['areas_of_development'], str) else obs['areas_of_development']
                        all_developments.extend(developments)
                    except:
                        pass

                if obs.get('recommendations'):
                    try:
                        recommendations = json.loads(obs['recommendations']) if isinstance(obs['recommendations'],
                                                                                           str) else obs[
                            'recommendations']
                        all_recommendations.extend(recommendations)
                    except:
                        pass

            # Calculate metrics for graphs
            total_observations = len(observations)
            active_goals = len([g for g in goal_progress if g.get('status') == 'active'])
            completed_goals = len([g for g in goal_progress if g.get('status') == 'achieved'])

            # Count frequency of strengths and development areas
            strength_counts = {}
            for strength in all_strengths:
                strength_counts[strength] = strength_counts.get(strength, 0) + 1

            development_counts = {}
            for dev in all_developments:
                development_counts[dev] = development_counts.get(dev, 0) + 1

            # Calculate weekly observation trends
            weekly_trends = self._calculate_weekly_trends(observations, year, month)
            learning_metrics = self._calculate_learning_metrics(observations)

            # Prepare graph data suggestions
            graph_suggestions = []

            if total_observations > 0:
                graph_suggestions.append({
                    "type": "line_chart",
                    "title": "📅 Weekly Observation Trends",
                    "description": f"Shows observation frequency across {len(weekly_trends)} weeks in the month",
                    "data": weekly_trends,
                    "xAxis": "Week",
                    "yAxis": "Number of Observations"
                })

                graph_suggestions.append({
                    "type": "bar_chart",
                    "title": "📚 Daily Learning Activities",
                    "description": f"Distribution of {total_observations} learning sessions throughout the month",
                    "data": {"total_sessions": total_observations},
                    "insights": f"Average of {total_observations / 30:.1f} sessions per day"
                })

            if strength_counts:
                graph_suggestions.append({
                    "type": "pie_chart",
                    "title": "🌟 Strength Areas Distribution",
                    "description": f"Breakdown of {len(strength_counts)} different strength categories",
                    "data": dict(list(strength_counts.items())[:8]),
                    "insights": f"Most frequent strength: {max(strength_counts.keys(), key=strength_counts.get)}"
                })

                graph_suggestions.append({
                    "type": "donut_chart",
                    "title": "🏆 Top 5 Strengths Focus",
                    "description": "Concentrated view of primary strength areas",
                    "data": dict(list(strength_counts.items())[:5])
                })

            if development_counts:
                graph_suggestions.append({
                    "type": "horizontal_bar",
                    "title": "📈 Development Priority Areas",
                    "description": f"Focus areas requiring attention with frequency analysis",
                    "data": dict(list(development_counts.items())[:6]),
                    "insights": f"Primary development focus: {max(development_counts.keys(), key=development_counts.get)}"
                })

            if goal_progress:
                goal_completion_rate = (completed_goals / (active_goals + completed_goals)) * 100 if (
                                                                                                             active_goals + completed_goals) > 0 else 0

                graph_suggestions.append({
                    "type": "gauge_chart",
                    "title": "🎯 Goal Achievement Rate",
                    "description": f"Monthly goal completion progress",
                    "data": {
                        "completion_rate": goal_completion_rate,
                        "completed": completed_goals,
                        "active": active_goals,
                        "total": active_goals + completed_goals
                    },
                    "insights": f"{goal_completion_rate:.1f}% completion rate this month"
                })

                graph_suggestions.append({
                    "type": "progress_bars",
                    "title": "📊 Individual Goal Progress",
                    "description": "Detailed progress tracking for each goal",
                    "data": [{"goal": g['goal_text'][:30], "progress": g['avg_score']} for g in goal_progress[:5]]
                })

            # Add learning engagement metrics if available
            if learning_metrics:
                graph_suggestions.append({
                    "type": "radar_chart",
                    "title": "📊 Learning Engagement Profile",
                    "description": "Multi-dimensional view of learning engagement",
                    "data": learning_metrics,
                    "insights": "Comprehensive engagement across different learning domains"
                })

            # Create comprehensive prompt for JSON generation with new fields
            monthly_prompt = f"""
            You are an AI assistant for a learning observation system. Generate a comprehensive monthly report based on the provided observation data for educational assessment and progress tracking.

            REPORTING PERIOD: {calendar.month_name[month]} {year}
            STUDENT: {child_name}
            TOTAL OBSERVATIONS: {total_observations}
            GOALS STATUS: {active_goals} active, {completed_goals} completed
            LEARNING SESSIONS ANALYZED: {len(observation_texts)}

            COMMUNICATION SKILLS SUMMARY: {json.dumps(comm_skills, indent=2)}
            GROWTH METRICS SUMMARY: {json.dumps(growth_metrics, indent=2)}

            OBSERVATION SAMPLE DATA: {json.dumps(observation_texts[:3], indent=2)}
            STRENGTHS IDENTIFIED: {list(strength_counts.keys())[:8]}
            DEVELOPMENT AREAS: {list(development_counts.keys())[:6]}
            WEEKLY TRENDS: {weekly_trends}

            QUANTIFIABLE METRICS FOR VISUAL ANALYTICS:
            {json.dumps(graph_suggestions, indent=2)}

            Format your response as JSON with the following structure:
            {{
              "studentName": "{child_name}",
              "studentId": "Monthly-{year}-{month:02d}-Report",
              "className": "📋 Monthly Learning Progress Assessment",
              "date": "{calendar.month_name[month]} {year}",
              "observations": "Comprehensive monthly learning summary combining all {total_observations} observation sessions. Detail the student's learning journey throughout {calendar.month_name[month]}, highlighting key educational milestones, skill development patterns, engagement levels, and notable learning breakthroughs. Include specific examples of learning activities, problem-solving approaches, creative expressions, and social interactions observed during the month.",
              "strengths": {json.dumps(list(strength_counts.keys())[:8])},
              "areasOfDevelopment": {json.dumps(list(development_counts.keys())[:6])},
              "recommendations": ["Specific actionable recommendations for {calendar.month_name[month + 1 if month < 12 else 1]} based on observed learning patterns", "Suggested learning activities to reinforce strengths", "Targeted interventions for development areas", "Parent engagement strategies", "Environmental modifications to support learning"],
              "communicationSkills": {json.dumps(comm_skills)},
              "growthMetrics": {json.dumps(growth_metrics)},
              "monthlyMetrics": {{
                "totalObservations": {total_observations},
                "activeGoals": {active_goals},
                "completedGoals": {completed_goals},
                "goalCompletionRate": {(completed_goals / (active_goals + completed_goals)) * 100 if (active_goals + completed_goals) > 0 else 0},
                "topStrengths": {dict(list(strength_counts.items())[:5])},
                "developmentFocus": {dict(list(development_counts.items())[:5])},
                "weeklyTrends": {weekly_trends},
                "averageSessionsPerWeek": {total_observations / 4.3 if total_observations > 0 else 0}
              }},
              "learningAnalytics": {{
                "engagement Level": "High/Medium/Low based on observation frequency and quality",
                "learning Velocity": "Assessment of learning pace and skill acquisition speed",
                "social Development": "Progress in social skills and peer interactions",
                "cognitive Growth": "Intellectual development and problem-solving abilities",
                "creativity Index": "Creative expression and innovative thinking patterns",
                "independence Level": "Self-directed learning and autonomous task completion"
              }},
              "suggestedGraphs": {graph_suggestions},
              "progressInsights": [
                "Key learning breakthroughs achieved this month",
                "Patterns in learning preferences and optimal learning conditions",
                "Social and emotional development observations",
                "Areas showing accelerated growth",
                "Challenges overcome and resilience demonstrated"
              ]
            }}

            For the observations field, provide a comprehensive narrative like:
            "Throughout {calendar.month_name[month]} {year}, {child_name} demonstrated remarkable growth across multiple learning domains. The student engaged in {total_observations} documented learning sessions, showing consistent curiosity and enthusiasm for discovery-based learning. Key highlights include [specific learning achievements], where the student mastered [specific skills] through hands-on exploration and guided inquiry. Notable progress was observed in [subject areas], with the student showing particular aptitude for [specific skills]. The learning journey included diverse activities such as [examples from observations], demonstrating the student's ability to connect concepts across different domains and apply learning in practical contexts."

            Include detailed quantifiable metrics and suggest appropriate visual analytics for comprehensive representation of the student's monthly learning progress. Be creative in extracting meaningful patterns, learning trajectories, and developmental insights from the observation data.

            Ensure all recommendations are specific, actionable, and tailored to the student's individual learning profile and developmental stage.
            """

            print("DEBUG: monthly_prompt", monthly_prompt)
            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                {"role": "user", "parts": [{"text": monthly_prompt}]}
            ])
            print("DEBUG: Gemini response", response.text)
            # Clean up Gemini's markdown code block if present
            cleaned = response.text.strip()
            if cleaned.startswith('```json'):
                cleaned = cleaned[len('```json'):].strip()
            if cleaned.startswith('```'):
                cleaned = cleaned[len('```'):].strip()
            if cleaned.endswith('```'):
                cleaned = cleaned[:-3].strip()
            # Validate JSON before returning
            try:
                json.loads(cleaned)
                return cleaned
            except Exception as e:
                print("DEBUG: Gemini did not return valid JSON.", str(e))
                return "Error: Gemini did not return valid JSON."

        except Exception as e:
            print("DEBUG: Exception in summary generation", str(e))
            return f"Error generating monthly summary: {str(e)}"

    def _calculate_weekly_trends(self, observations, year, month):
        """Calculate weekly observation trends for the month"""
        weekly_counts = [0, 0, 0, 0, 0]  # Up to 5 weeks in a month

        for obs in observations:
            try:
                obs_date = datetime.strptime(obs.get('date', ''), '%Y-%m-%d')
                # Calculate which week of the month
                week_of_month = (obs_date.day - 1) // 7
                if week_of_month < 5:
                    weekly_counts[week_of_month] += 1
            except:
                continue

        return {f"Week {i + 1}": count for i, count in enumerate(weekly_counts) if count > 0}

    def _calculate_learning_metrics(self, observations):
        """Calculate learning engagement and progress metrics"""
        if not observations:
            return {}

        # Calculate various learning metrics
        total_sessions = len(observations)

        # Analyze themes and curiosity seeds
        themes = []
        curiosity_seeds = []

        for obs in observations:
            if obs.get('theme_of_day'):
                themes.append(obs['theme_of_day'])
            if obs.get('curiosity_seed'):
                curiosity_seeds.append(obs['curiosity_seed'])

        return {
            "session_frequency": total_sessions,
            "theme_diversity": len(set(themes)),
            "curiosity_engagement": len(set(curiosity_seeds)),
            "learning_consistency": "High" if total_sessions > 15 else "Medium" if total_sessions > 8 else "Low"
        }

    def generate_excel_report(self, observations, goal_progress, strength_counts, development_counts):
        """Generate Excel report with multiple sheets, including curiosity and growth line charts"""
        buffer = io.BytesIO()

        # --- Extract curiosity and growth scores by date ---
        curiosity_by_date = {}
        growth_by_date = {}
        comm_skills_by_date = {}
        for obs in observations:
            date = obs.get('date')
            try:
                full_data = json.loads(obs.get('full_data', '{}'))
                report = full_data.get('formatted_report', '')
            except Exception:
                report = ''
            # Extract curiosity score
            curiosity_match = re.search(r'🌈 Curiosity Response Index: (\d{1,2}) ?/ ?10', report)
            if curiosity_match:
                curiosity_score = int(curiosity_match.group(1))
                curiosity_by_date[date] = curiosity_score
            # Extract growth score (X/7)
            growth_match = re.search(r'Overall Growth Score.*?(\d)\s*/\s*7', report)
            if growth_match:
                growth_score = int(growth_match.group(1))
                growth_by_date[date] = growth_score
            # Extract communication skills
            comm_skills = {}
            confidence_match = re.search(r'Confidence level:.*?(Strong|Moderate|Weak)', report, re.IGNORECASE)
            if confidence_match:
                comm_skills['confidence'] = confidence_match.group(1)
            clarity_match = re.search(r'Clarity of thought:.*?(Clear|Moderate|Unclear)', report, re.IGNORECASE)
            if clarity_match:
                comm_skills['clarity'] = clarity_match.group(1)
            participation_match = re.search(r'Participation.*?:.*?(Active|Moderate|Passive)', report, re.IGNORECASE)
            if participation_match:
                comm_skills['participation'] = participation_match.group(1)
            sequence_match = re.search(r'Sequence of explanation:.*?(Logical|Moderate|Disorganized)', report,
                                       re.IGNORECASE)
            if sequence_match:
                comm_skills['sequencing'] = sequence_match.group(1)
            if comm_skills:
                comm_skills_by_date[date] = comm_skills

        # Sort by date
        curiosity_dates = sorted(curiosity_by_date.keys())
        growth_dates = sorted(growth_by_date.keys())
        comm_dates = sorted(comm_skills_by_date.keys())
        curiosity_scores = [curiosity_by_date[d] for d in curiosity_dates]
        growth_scores = [growth_by_date[d] for d in growth_dates]
        comm_skills_data = [comm_skills_by_date[d] for d in comm_dates]

        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            workbook = writer.book
            # --- Summary sheet ---
            summary_data = {
                "Metric": ["Total Observations", "Goals Tracked", "Average Goal Score"],
                "Value": [len(observations), len(goal_progress),
                          sum(g['avg_score'] for g in goal_progress) / len(goal_progress) if goal_progress else 0]
            }
            summary_df = pd.DataFrame(summary_data)
            summary_df.to_excel(writer, sheet_name='Summary', index=False)
            worksheet = writer.sheets['Summary']
            # Add a narrative summary at the top
            worksheet.write('A1', '📋 Monthly Learning Progress Summary',
                            workbook.add_format({'bold': True, 'font_size': 14}))
            worksheet.write('A3', '📊 Key Metrics:', workbook.add_format({'bold': True}))

            # --- Curiosity Line Chart ---
            if curiosity_dates:
                curiosity_df = pd.DataFrame({
                    'Date': curiosity_dates,
                    'Curiosity Score': curiosity_scores
                })
                curiosity_df.to_excel(writer, sheet_name='Curiosity Trend', index=False)
                ws = writer.sheets['Curiosity Trend']
                chart = workbook.add_chart({'type': 'line'})
                chart.add_series({
                    'name': 'Curiosity Score',
                    'categories': ['Curiosity Trend', 1, 0, len(curiosity_df), 0],
                    'values': ['Curiosity Trend', 1, 1, len(curiosity_df), 1],
                })
                chart.set_title({'name': '🌈 Curiosity Response Index by Day'})
                chart.set_x_axis({'name': 'Date'})
                chart.set_y_axis({'name': 'Curiosity Score', 'major_gridlines': {'visible': False}})
                ws.insert_chart('D2', chart)

            # --- Growth Line Chart ---
            if growth_dates:
                growth_df = pd.DataFrame({
                    'Date': growth_dates,
                    'Growth Score': growth_scores
                })
                growth_df.to_excel(writer, sheet_name='Growth Trend', index=False)
                ws = writer.sheets['Growth Trend']
                chart = workbook.add_chart({'type': 'line'})
                chart.add_series({
                    'name': 'Growth Score',
                    'categories': ['Growth Trend', 1, 0, len(growth_df), 0],
                    'values': ['Growth Trend', 1, 1, len(growth_df), 1],
                })
                chart.set_title({'name': '📈 Overall Growth Score by Day'})
                chart.set_x_axis({'name': 'Date'})
                chart.set_y_axis({'name': 'Growth Score', 'major_gridlines': {'visible': False}})
                ws.insert_chart('D2', chart)

            # --- Communication Skills Sheet ---
            if comm_dates:
                comm_df = pd.DataFrame({
                    'Date': comm_dates,
                    'Confidence': [d.get('confidence', '') for d in comm_skills_data],
                    'Clarity': [d.get('clarity', '') for d in comm_skills_data],
                    'Participation': [d.get('participation', '') for d in comm_skills_data],
                    'Sequencing': [d.get('sequencing', '') for d in comm_skills_data]
                })
                comm_df.to_excel(writer, sheet_name='Communication Skills', index=False)

            # --- Strengths sheet ---
            if strength_counts:
                strengths_df = pd.DataFrame([
                    {"Strength": strength, "Count": count}
                    for strength, count in strength_counts.items()
                ])
                strengths_df.to_excel(writer, sheet_name='🌟 Strengths', index=False)

            # --- Development areas sheet ---
            if development_counts:
                development_df = pd.DataFrame([
                    {"Development Area": area, "Count": count}
                    for area, count in development_counts.items()
                ])
                development_df.to_excel(writer, sheet_name='📈 Development Areas', index=False)

            # --- Goal progress sheet ---
            if goal_progress:
                goals_df = pd.DataFrame([
                    {"Goal": g['goal_text'], "Average Score": g['avg_score'], "Observations": g['num_observations']}
                    for g in goal_progress
                ])
                goals_df.to_excel(writer, sheet_name='🎯 Goal Progress', index=False)

        buffer.seek(0)
        return buffer

    def generate_monthly_docx_report(self, observations, goal_progress, strength_counts, development_counts,
                                     summary_json):
        doc = docx.Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Segoe UI'
        font.size = Pt(11)

        # --- Narrative Section ---
        doc.add_heading(f"📋 Monthly Growth Report", 0)
        doc.add_paragraph(f"📅 {summary_json.get('date', '')}")
        doc.add_paragraph(f"👧 Student: {summary_json.get('studentName', '')}")
        doc.add_paragraph("")
        doc.add_paragraph(summary_json.get('observations', ''))
        doc.add_paragraph("")

        # --- Strengths ---
        doc.add_heading("🌟 Strengths Observed", level=1)
        for s in summary_json.get('strengths', []):
            doc.add_paragraph(s, style='List Bullet')
        doc.add_paragraph("")

        # --- Areas for Development ---
        doc.add_heading("📈 Areas for Development", level=1)
        for a in summary_json.get('areasOfDevelopment', []):
            doc.add_paragraph(a, style='List Bullet')
        doc.add_paragraph("")

        # --- Communication Skills ---
        doc.add_heading("🗣️ Communication Skills", level=1)
        comm_skills = summary_json.get('communicationSkills', {})
        if comm_skills:
            doc.add_paragraph(f"Confidence level: {comm_skills.get('confidence', 'No data').title()}")
            doc.add_paragraph(f"Clarity of thought: {comm_skills.get('clarity', 'No data').title()}")
            doc.add_paragraph(f"Participation & engagement: {comm_skills.get('participation', 'No data').title()}")
            doc.add_paragraph(f"Sequence of explanation: {comm_skills.get('sequencing', 'No data').title()}")
        doc.add_paragraph("")

        # --- Growth Metrics ---
        doc.add_heading("📊 Growth Metrics", level=1)
        growth_metrics = summary_json.get('growthMetrics', {})
        if growth_metrics:
            for area, rating in growth_metrics.items():
                doc.add_paragraph(f"{area}: {rating.title()}")
        doc.add_paragraph("")

        # --- Recommendations ---
        doc.add_heading("📣 Recommendations for Next Month", level=1)
        for r in summary_json.get('recommendations', []):
            doc.add_paragraph(r, style='List Bullet')
        doc.add_paragraph("")

        # --- Learning Analytics ---
        doc.add_heading("📊 Learning Analytics", level=1)
        analytics = summary_json.get('learningAnalytics', {})
        for k, v in analytics.items():
            doc.add_paragraph(f"{k.replace('_', ' ').title()}: {v}")
        doc.add_paragraph("")

        # --- Progress Insights ---
        doc.add_heading("🔍 Progress Insights", level=1)
        for insight in summary_json.get('progressInsights', []):
            doc.add_paragraph(insight, style='List Bullet')
        doc.add_paragraph("")

        # --- Graphs Section ---
        doc.add_heading("📈 Visual Analytics", level=1)

        # Curiosity and Growth scores by day (parse from observations)
        curiosity_by_date = {}
        growth_by_date = {}
        for obs in observations:
            date = obs.get('date')
            try:
                full_data = json.loads(obs.get('full_data', '{}'))
                report = full_data.get('formatted_report', '')
            except Exception:
                report = ''
            curiosity_match = re.search(r'🌈 Curiosity Response Index: (\d{1,2}) ?/ ?10', report)
            if curiosity_match:
                curiosity_score = int(curiosity_match.group(1))
                curiosity_by_date[date] = curiosity_score
            growth_match = re.search(r'Overall Growth Score.*?(\d)\s*/\s*7', report)
            if growth_match:
                growth_score = int(growth_match.group(1))
                growth_by_date[date] = growth_score
        # Sort by date
        curiosity_dates = sorted(curiosity_by_date.keys())
        growth_dates = sorted(growth_by_date.keys())
        curiosity_scores = [curiosity_by_date[d] for d in curiosity_dates]
        growth_scores = [growth_by_date[d] for d in growth_dates]

        # --- Curiosity Line Chart ---
        if curiosity_dates:
            fig, ax = plt.subplots()
            ax.plot(curiosity_dates, curiosity_scores, marker='o', color='blue')
            ax.set_title('🌈 Curiosity Response Index by Day')
            ax.set_xlabel('Date')
            ax.set_ylabel('Curiosity Score')
            ax.xaxis.set_major_locator(MaxNLocator(integer=True))
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            img_stream = BytesIO()
            plt.savefig(img_stream, format='png')
            plt.close(fig)
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.add_paragraph("")

        # --- Growth Line Chart ---
        if growth_dates:
            fig, ax = plt.subplots()
            ax.plot(growth_dates, growth_scores, marker='o', color='green')
            ax.set_title('📈 Overall Growth Score by Day')
            ax.set_xlabel('Date')
            ax.set_ylabel('Growth Score')
            ax.xaxis.set_major_locator(MaxNLocator(integer=True))
            plt.xticks(rotation=45, ha='right')
            plt.tight_layout()
            img_stream = BytesIO()
            plt.savefig(img_stream, format='png')
            plt.close(fig)
            img_stream.seek(0)
            doc.add_picture(img_stream, width=Inches(5.5))
            doc.add_paragraph("")

        # --- Other suggested graphs from summary_json ---
        for graph in summary_json.get('suggestedGraphs', []):
            if graph['type'] in ['line_chart', 'bar_chart']:
                fig, ax = plt.subplots()
                if graph['type'] == 'line_chart':
                    x = list(graph['data'].keys())
                    y = list(graph['data'].values())
                    ax.plot(x, y, marker='o')
                elif graph['type'] == 'bar_chart':
                    x = list(graph['data'].keys())
                    y = list(graph['data'].values())
                    ax.bar(x, y)
                ax.set_title(graph.get('title', ''))
                ax.set_xlabel(graph.get('xAxis', ''))
                ax.set_ylabel(graph.get('yAxis', ''))
                plt.xticks(rotation=45, ha='right')
                plt.tight_layout()
                img_stream = BytesIO()
                plt.savefig(img_stream, format='png')
                plt.close(fig)
                img_stream.seek(0)
                doc.add_picture(img_stream, width=Inches(5.5))
                doc.add_paragraph(graph.get('description', ''))
                doc.add_paragraph("")

        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        return docx_bytes

    def generate_monthly_pdf_report(self, observations, goal_progress, strength_counts, development_counts,
                                    summary_json):

        import tempfile
        import os
        docx_bytes = self.generate_monthly_docx_report(observations, goal_progress, strength_counts, development_counts,
                                                       summary_json)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx.write(docx_bytes.read())
            tmp_docx_path = tmp_docx.name
        tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
        convert(tmp_docx_path, tmp_pdf_path)
        with open(tmp_pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        os.remove(tmp_docx_path)
        os.remove(tmp_pdf_path)
        from io import BytesIO
        return BytesIO(pdf_bytes)
